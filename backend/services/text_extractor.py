# backend/services/text_extractor.py

from __future__ import annotations

import io
import logging
import os
import shutil
from enum import Enum
from typing import Optional

import pdfplumber
import docx

logger = logging.getLogger(__name__)


# -----------------------------
# Configuration (env overridable)
# -----------------------------
class Config:
    MAX_CHARS: int = int(os.getenv("EXTRACTOR_MAX_CHARS", "80000"))
    OCR_MAX_PAGES: int = int(os.getenv("OCR_MAX_PAGES", "12"))
    OCR_DPI: int = int(os.getenv("OCR_DPI", "200"))
    OCR_LANG: str = os.getenv("OCR_LANG", "eng")
    OCR_ENABLED: bool = os.getenv("OCR_ENABLED", "true").lower() == "true"

    # NOTE: file-size limit should ideally be enforced at router level,
    # but we keep a safety guard here too.
    MAX_FILE_SIZE_MB: int = int(os.getenv("MAX_FILE_SIZE_MB", "50"))

    @classmethod
    def max_file_bytes(cls) -> int:
        return cls.MAX_FILE_SIZE_MB * 1024 * 1024


# -----------------------------
# Exceptions (clearer debugging)
# -----------------------------
class ExtractionError(Exception):
    pass


class UnsupportedFormatError(ExtractionError):
    pass


class EmptyContentError(ExtractionError):
    pass


class FileTooLargeError(ExtractionError):
    pass


class OCRError(ExtractionError):
    pass


# -----------------------------
# File type detection
# -----------------------------
class FileType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    UNKNOWN = "unknown"


def detect_file_type(content_type: Optional[str], filename: Optional[str]) -> FileType:
    filename_lower = (filename or "").lower()
    ct = (content_type or "").lower()

    if ct == "application/pdf" or filename_lower.endswith(".pdf"):
        return FileType.PDF

    if (
        ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename_lower.endswith(".docx")
    ):
        return FileType.DOCX

    if ct == "application/msword" or filename_lower.endswith(".doc"):
        return FileType.DOC

    if ct == "text/plain" or filename_lower.endswith(".txt"):
        return FileType.TXT

    return FileType.UNKNOWN


# -----------------------------
# Tesseract setup (Windows + Linux safe)
# -----------------------------
def _resolve_tesseract_path() -> Optional[str]:
    candidates = [
        os.getenv("TESSERACT_CMD"),
        shutil.which("tesseract"),
        # common Windows paths
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        # common linux paths
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
    ]

    for path in candidates:
        if path and os.path.isfile(path):
            return path
    return None


def _configure_tesseract() -> bool:
    if not Config.OCR_ENABLED:
        logger.info("OCR disabled via OCR_ENABLED=false")
        return False

    try:
        import pytesseract  # noqa
    except Exception:
        logger.warning("pytesseract not installed; OCR disabled")
        return False

    tpath = _resolve_tesseract_path()
    if not tpath:
        logger.warning("tesseract binary not found; OCR disabled")
        return False

    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = tpath
    logger.info(f"OCR enabled. Using tesseract at: {tpath}")
    return True


_TESSERACT_AVAILABLE = _configure_tesseract()


# -----------------------------
# PDF Extraction (native + OCR fallback)
# -----------------------------
def _extract_pdf_native(file_bytes: bytes) -> str:
    parts: list[str] = []
    total_chars = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            try:
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    parts.append(page_text)
                    total_chars += len(page_text)

                if total_chars >= Config.MAX_CHARS:
                    logger.info(f"Reached MAX_CHARS during native PDF extraction at page {idx}")
                    break
            except Exception as exc:
                logger.warning(f"Native extract failed on page {idx}: {exc}")
                continue

    return "\n\n".join(parts).strip()


def _extract_pdf_ocr(file_bytes: bytes) -> str:
    if not _TESSERACT_AVAILABLE:
        raise OCRError("OCR not available (tesseract/pytesseract missing or disabled)")

    try:
        import pytesseract
        from PIL import Image  # noqa: F401
    except Exception as exc:
        raise OCRError("OCR requires Pillow + pytesseract installed") from exc

    parts: list[str] = []
    total_chars = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        total_pages = len(pdf.pages)
        pages_to_process = min(total_pages, Config.OCR_MAX_PAGES)

        logger.info(f"OCR fallback: processing {pages_to_process}/{total_pages} pages")

        for idx in range(pages_to_process):
            page_num = idx + 1
            try:
                page = pdf.pages[idx]
                pil_image = page.to_image(resolution=Config.OCR_DPI).original

                # Keep OCR config minimal + stable (avoid risky OSD changes)
                page_text = pytesseract.image_to_string(pil_image, lang=Config.OCR_LANG)
                page_text = (page_text or "").strip()

                if page_text:
                    parts.append(page_text)
                    total_chars += len(page_text)

                if total_chars >= Config.MAX_CHARS:
                    logger.info(f"Reached MAX_CHARS during OCR at page {page_num}")
                    break
            except Exception as exc:
                logger.warning(f"OCR failed on page {page_num}: {exc}")
                continue

    return "\n\n".join(parts).strip()


def extract_from_pdf(file_bytes: bytes, allow_ocr: bool = True) -> str:
    text = _extract_pdf_native(file_bytes)
    if text.strip():
        return text

    if allow_ocr:
        logger.info("No selectable text found; attempting OCR fallback")
        text = _extract_pdf_ocr(file_bytes)
        if text.strip():
            return text

    raise EmptyContentError("PDF contains no extractable text (and OCR failed or disabled)")


# -----------------------------
# DOCX Extraction (paragraphs + tables)
# -----------------------------
def extract_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ExtractionError(f"DOCX parsing failed: {exc}") from exc

    parts: list[str] = []
    total_chars = 0

    # paragraphs
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
            total_chars += len(t)
            if total_chars >= Config.MAX_CHARS:
                break

    # tables (only if still under cap)
    if total_chars < Config.MAX_CHARS:
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
                if row_text:
                    parts.append(row_text)
                    total_chars += len(row_text)
                    if total_chars >= Config.MAX_CHARS:
                        break
            if total_chars >= Config.MAX_CHARS:
                break

    text = "\n\n".join(parts).strip()
    if not text:
        raise EmptyContentError("DOCX contains no extractable text")
    return text


# -----------------------------
# TXT Extraction (multi-encoding)
# -----------------------------
def extract_from_txt(file_bytes: bytes) -> str:
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            t = file_bytes.decode(enc).strip()
            if t:
                return t
        except Exception:
            continue
    raise EmptyContentError("Could not decode TXT file using known encodings")


# -----------------------------
# Public API (backward compatible)
# -----------------------------
def extract_text(
    file_bytes: bytes,
    content_type: Optional[str] = None,
    filename: Optional[str] = None,
    allow_ocr: bool = True,
) -> str:
    """
    Unified extractor used by summarizer and other tools.

    Returns:
        Extracted text (capped at MAX_CHARS)

    Raises:
        FileTooLargeError, UnsupportedFormatError, EmptyContentError, ExtractionError
    """
    if not file_bytes:
        raise EmptyContentError("Empty file bytes")

    # Safety guard (router should also enforce)
    if len(file_bytes) > Config.max_file_bytes():
        raise FileTooLargeError(
            f"File too large: {len(file_bytes) / 1024 / 1024:.1f}MB (limit {Config.MAX_FILE_SIZE_MB}MB)"
        )

    ft = detect_file_type(content_type, filename)
    logger.info(f"Extracting text | type={ft.value} | filename={filename or 'unnamed'}")

    if ft == FileType.PDF:
        text = extract_from_pdf(file_bytes, allow_ocr=allow_ocr)
    elif ft in (FileType.DOCX, FileType.DOC):
        if ft == FileType.DOC:
            logger.warning("Legacy .doc detected; extraction may be incomplete (treated as DOCX parser)")
        text = extract_from_docx(file_bytes)
    elif ft == FileType.TXT:
        text = extract_from_txt(file_bytes)
    else:
        raise UnsupportedFormatError(
            f"Unsupported format. Got content_type={content_type}, filename={filename}"
        )

    text = (text or "").strip()
    if not text:
        raise EmptyContentError("No extractable text found (PDF may be scanned; OCR required)")

    if len(text) > Config.MAX_CHARS:
        logger.info(f"Truncating extracted text from {len(text)} to {Config.MAX_CHARS} chars")
        text = text[:Config.MAX_CHARS]

    return text


def is_ocr_available() -> bool:
    return _TESSERACT_AVAILABLE
